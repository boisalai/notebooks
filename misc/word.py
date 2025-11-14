"""
Markdown to Word (DOCX) Converter

A professional tool for converting Markdown documents to Word format with advanced
formatting capabilities including custom styles, headers, footers, images, tables,
and footnotes.

Dependencies:
    uv add python-docx pandoc

Requirements:
    - python-docx: Word document manipulation
    - pandoc: Document conversion (external dependency - brew install pandoc)
"""

from typing import Optional, Dict, List, Tuple
from enum import Enum
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from docx.table import _Cell
import subprocess
import shutil
import re
import logging
import platform


# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(levelname)s: %(message)s'
)
logger = logging.getLogger(__name__)


class PaperSize(Enum):
    """Supported paper sizes for document layout."""
    LETTER = "letter"  # 8.5 x 11 inches (215.9 x 279.4 mm)
    LEGAL = "legal"    # 8.5 x 14 inches (215.9 x 355.6 mm)
    A4 = "a4"          # 8.27 x 11.69 inches (210 x 297 mm)


class DocumentStyle(Enum):
    """Predefined document style templates."""
    REPORT = "report"
    NOTE = "note"
    LETTER = "letter"
    MEMO = "memo"


class DocumentConfig:
    """
    Configuration for document conversion and styling.

    Attributes:
        style: Document style template
        paper_size: Paper size for the document
        author: Document author name
        date: Document date
        heading_colors: RGB color tuples for heading levels 1-3
        footer_text: Footer text for odd and even pages
        font_name: Base font family
        base_font_size: Base font size in points
        margins: Page margins (top, right, bottom, left) in cm
        line_spacing: Line spacing multiplier
        generate_toc: Whether to generate table of contents
        language: Document language code (e.g., 'en-US', 'fr-CA')
        center_title: Whether to center the main title
    """

    # Style constants
    DEFAULT_TITLE_SIZE = 24
    DEFAULT_HEADING_1_SIZE = 18
    DEFAULT_HEADING_2_SIZE = 16
    DEFAULT_HEADING_3_SIZE = 14
    DEFAULT_TABLE_FONT_SIZE = 10
    DEFAULT_FOOTER_FONT_SIZE = 10
    DEFAULT_FOOTNOTE_FONT_SIZE = 10
    DEFAULT_PARAGRAPH_SPACING = 6
    DEFAULT_TABLE_CELL_SPACING = 2

    def __init__(
        self,
        style: DocumentStyle = DocumentStyle.REPORT,
        paper_size: PaperSize = PaperSize.LETTER,
        author: str = "",
        date: str = "",
        heading_colors: Dict[int, Tuple[int, int, int]] = None,
        footer_text: Dict[str, str] = None,
        font_name: str = "Arial",
        base_font_size: int = 12,
        margins: Tuple[float, float, float, float] = (2, 2, 2, 2),
        line_spacing: float = 1.0,
        generate_toc: bool = True,
        language: str = "en-US",
        center_title: bool = True
    ):
        self.style = style
        self.paper_size = paper_size
        self.author = author
        self.date = date
        self.heading_colors = heading_colors or {
            1: (37, 150, 190),
            2: (37, 150, 190),
            3: (37, 150, 190)
        }
        self.footer_text = footer_text or {
            "odd": "Page",
            "even": "Page"
        }
        self.font_name = font_name
        self.base_font_size = base_font_size
        self.margins = margins
        self.line_spacing = line_spacing
        self.generate_toc = generate_toc
        self.language = language
        self.center_title = center_title

        # Validate configuration
        self._validate()

    def _validate(self) -> None:
        """Validate configuration parameters."""
        if self.base_font_size <= 0:
            raise ValueError(f"Base font size must be positive, got {self.base_font_size}")

        if any(m < 0 for m in self.margins):
            raise ValueError(f"Margins must be non-negative, got {self.margins}")

        if self.line_spacing <= 0:
            raise ValueError(f"Line spacing must be positive, got {self.line_spacing}")

        for level, color in self.heading_colors.items():
            if not isinstance(color, tuple) or len(color) != 3:
                raise ValueError(f"Color for heading {level} must be RGB tuple, got {color}")
            if not all(0 <= c <= 255 for c in color):
                raise ValueError(f"RGB values must be 0-255, got {color}")

        if not isinstance(self.footer_text, dict) or 'odd' not in self.footer_text or 'even' not in self.footer_text:
            raise ValueError("footer_text must be a dict with 'odd' and 'even' keys")

    @classmethod
    def create_report_style(cls, **kwargs):
        """Create a preconfigured style for reports."""
        default_config = {
            'style': DocumentStyle.REPORT,
            'paper_size': PaperSize.LETTER,
            'heading_colors': {
                1: (37, 150, 190),
                2: (37, 150, 190),
                3: (37, 150, 190)
            },
            'footer_text': {
                "odd": "Right text | Page",
                "even": "Page | Left text"
            },
            'center_title': True
        }
        default_config.update(kwargs)
        return cls(**default_config)

    @classmethod
    def create_note_style(cls, **kwargs):
        """Create a preconfigured style for internal notes."""
        default_config = {
            'style': DocumentStyle.NOTE,
            'paper_size': PaperSize.LEGAL,
            'heading_colors': {
                1: (70, 70, 70),
                2: (100, 100, 100),
                3: (130, 130, 130)
            },
            'footer_text': {
                "odd": "Internal Note | Page",
                "even": "Page | Internal Note"
            },
            'margins': (1.5, 1.5, 1.5, 1.5),
            'center_title': True
        }
        default_config.update(kwargs)
        return cls(**default_config)


class MarkdownToDocxConverter:
    """
    Convert Markdown documents to Word (DOCX) format with advanced formatting.

    This converter uses Pandoc for initial conversion and applies extensive
    post-processing to ensure professional formatting including:
    - Custom heading styles and colors
    - Image embedding with size adjustment
    - Table formatting with borders
    - Footnote styling
    - Custom headers and footers with odd/even page support
    - Language settings
    - Centered title option
    """

    # Image processing constants
    MAX_IMAGE_WIDTH = Inches(6)
    SUPPORTED_IMAGE_EXTENSIONS = {'.png', '.jpg', '.jpeg', '.gif', '.bmp'}

    # Compiled regex patterns
    IMAGE_PATTERN = re.compile(r'!\[(.*?)\]\((.*?)\)')

    def __init__(self, config: DocumentConfig = None, verbose: bool = True):
        """
        Initialize the converter.

        Args:
            config: Document configuration. If None, uses default settings.
            verbose: Enable verbose logging output.
        """
        self.config = config or DocumentConfig()
        self.verbose = verbose

        if not verbose:
            logger.setLevel(logging.WARNING)

        self._check_dependencies()

    def _check_dependencies(self) -> None:
        """Verify that Pandoc is installed on the system."""
        if not shutil.which('pandoc'):
            system = platform.system()
            install_instructions = {
                'Darwin': 'brew install pandoc',
                'Linux': 'sudo apt-get install pandoc  # or use your package manager',
                'Windows': 'Download from https://pandoc.org/installing.html'
            }
            instruction = install_instructions.get(system, 'See https://pandoc.org/installing.html')

            raise RuntimeError(
                f"Pandoc is not installed.\n"
                f"Installation: {instruction}"
            )

    def _extract_title_from_markdown(self, content: str) -> str:
        """
        Extract the main title (H1) from markdown content.

        Args:
            content: Markdown document content

        Returns:
            The first H1 heading found, or "Untitled Document" if none exists
        """
        lines = content.split('\n')
        for line in lines:
            if line.startswith('# '):
                return line[2:].strip()
        return "Untitled Document"

    def _extract_image_references(self, content: str) -> List[dict]:
        """
        Extract image references from markdown content.

        Args:
            content: Markdown document content

        Returns:
            List of dictionaries containing image metadata
        """
        matches = self.IMAGE_PATTERN.findall(content)

        image_refs = []
        for alt_text, path in matches:
            path = path.strip()
            # Remove 'img/' prefix if present
            if path.startswith('img/'):
                path = path[4:]

            image_refs.append({
                'alt_text': alt_text,
                'path': path,
                'original_markdown': f'![{alt_text}]({path})'
            })

        logger.info(f"Found {len(image_refs)} image references")
        return image_refs

    def _remove_image_references(self, content: str) -> str:
        """
        Replace image references with placeholders.

        Args:
            content: Markdown document content

        Returns:
            Content with images replaced by placeholders
        """
        return self.IMAGE_PATTERN.sub(r'[IMAGE_PLACEHOLDER]', content)

    def _run_pandoc_conversion(self, input_path: Path, output_path: Path, title: str) -> None:
        """
        Run pandoc conversion with custom heading level mapping.

        This method uses a Lua filter to adjust heading levels so that
        Markdown H2 becomes Word Heading 1, H3 becomes Heading 2, etc.

        Args:
            input_path: Path to input markdown file
            output_path: Path to output docx file
            title: Document title
        """
        lua_path = None
        try:
            lua_path = self._create_lua_script(input_path)
            cmd = self._build_pandoc_command(input_path, output_path, title, lua_path)

            result = subprocess.run(cmd, check=True, capture_output=True, text=True)
            logger.info("Pandoc conversion completed successfully")

        except subprocess.CalledProcessError as e:
            logger.error(f"Pandoc conversion failed: {e.stderr}")
            raise RuntimeError(f"Pandoc conversion failed: {e.stderr}")
        except Exception as e:
            logger.error(f"Error during pandoc conversion: {str(e)}")
            raise
        finally:
            if lua_path:
                self._cleanup_lua_script(lua_path)

    def _create_lua_script(self, input_path: Path) -> Path:
        """
        Create a temporary Lua script for pandoc filtering.

        The script adjusts heading levels to match the desired hierarchy.

        Args:
            input_path: Path to input file (used to determine script location)

        Returns:
            Path to the created Lua script
        """
        lua_script = """
        function Header(el)
            if el.level > 1 then
                el.level = el.level - 1
            end
            return el
        end
        """
        lua_path = input_path.parent / "adjust_headers.lua"

        if lua_path.exists():
            logger.warning(f"Lua script already exists at {lua_path}, overwriting")

        with open(lua_path, 'w', encoding='utf-8') as f:
            f.write(lua_script)
        return lua_path

    def _build_pandoc_command(self, input_path: Path, output_path: Path,
                             title: str, lua_path: Path) -> List[str]:
        """
        Build the pandoc command with all necessary options.

        Args:
            input_path: Input markdown file
            output_path: Output docx file
            title: Document title
            lua_path: Path to Lua filter script

        Returns:
            Command as list of strings
        """
        cmd = [
            'pandoc',
            str(input_path),
            '-o', str(output_path),
            '-f', 'markdown',
            '-t', 'docx',
            '--wrap=none',
            '--columns=999',
            '--lua-filter=' + str(lua_path),
            '-M', f'title={title}',
            '-M', f'author={self.config.author}',
            '-M', f'date={self.config.date}'
        ]

        if self.config.generate_toc:
            cmd.extend(['--toc', '--number-sections'])

        return cmd

    def _cleanup_lua_script(self, lua_path: Path) -> None:
        """
        Remove temporary Lua script.

        Args:
            lua_path: Path to Lua script to delete
        """
        try:
            if lua_path.exists():
                lua_path.unlink()
        except Exception as e:
            logger.warning(f"Could not delete temporary Lua file: {e}")

    def convert(self,
                input_file: str,
                output_file: str,
                working_dir: Optional[str] = None) -> None:
        """
        Convert a Markdown file to Word format.

        Args:
            input_file: Name of the input markdown file
            output_file: Name for the output docx file
            working_dir: Working directory (defaults to current directory)

        Raises:
            FileNotFoundError: If input file doesn't exist
            RuntimeError: If conversion fails
        """
        work_dir, input_path, output_path = self._setup_paths(
            input_file, output_file, working_dir
        )

        if not input_path.exists():
            raise FileNotFoundError(f"Input file not found: {input_path}")

        if input_path.suffix.lower() not in {'.md', '.markdown'}:
            logger.warning(f"Input file {input_path} may not be a markdown file")

        img_dir = self._create_image_directory(work_dir)
        content = self._read_markdown_content(input_path)
        document_title = self._extract_title_from_markdown(content)
        image_refs = self._extract_image_references(content)
        temp_md = self._create_temp_markdown(content, work_dir, input_path)

        try:
            self._run_pandoc_conversion(temp_md, output_path, document_title)
            self._post_process_document(output_path, document_title, image_refs, work_dir)
            logger.info(f"Conversion successful! File saved: {output_path}")
        finally:
            self._cleanup_temp_markdown(temp_md)

    def _setup_paths(self, input_file: str, output_file: str,
                     working_dir: Optional[str]) -> Tuple[Path, Path, Path]:
        """
        Setup and validate file paths.

        Args:
            input_file: Input filename or absolute path
            output_file: Output filename or absolute path
            working_dir: Working directory (used for relative paths and temp files)

        Returns:
            Tuple of (working_dir, input_path, output_path)
        """
        work_dir = Path(working_dir).resolve() if working_dir else Path.cwd()

        if not work_dir.exists():
            raise FileNotFoundError(f"Working directory not found: {work_dir}")

        input_path = Path(input_file)
        if not input_path.is_absolute():
            input_path = work_dir / input_file
        input_path = input_path.resolve()

        output_path = Path(output_file)
        if not output_path.is_absolute():
            output_path = work_dir / output_file
        output_path = output_path.resolve()

        return work_dir, input_path, output_path

    def _create_image_directory(self, work_dir: Path) -> Path:
        """
        Create image directory if it doesn't exist.

        Args:
            work_dir: Working directory

        Returns:
            Path to image directory
        """
        img_dir = work_dir / "img"
        if not img_dir.exists():
            img_dir.mkdir(parents=True)
            logger.info(f"Created image directory: {img_dir}")
        return img_dir

    def _read_markdown_content(self, input_path: Path) -> str:
        """
        Read markdown file content.

        Args:
            input_path: Path to markdown file

        Returns:
            File content as string

        Raises:
            IOError: If file cannot be read
        """
        try:
            with open(input_path, 'r', encoding='utf-8') as f:
                return f.read()
        except IOError as e:
            raise IOError(f"Cannot read markdown file: {e}")

    def _create_temp_markdown(self, content: str, work_dir: Path,
                             input_path: Path) -> Path:
        """
        Create temporary markdown file without image references.

        Args:
            content: Original markdown content
            work_dir: Working directory
            input_path: Original input path

        Returns:
            Path to temporary file
        """
        temp_content = self._remove_image_references(content)
        temp_md = work_dir / f"temp_{input_path.name}"

        try:
            with open(temp_md, 'w', encoding='utf-8') as f:
                f.write(temp_content)
        except IOError as e:
            raise IOError(f"Cannot create temporary markdown file: {e}")

        return temp_md

    def _cleanup_temp_markdown(self, temp_md: Path) -> None:
        """
        Remove temporary markdown file.

        Args:
            temp_md: Path to temporary file
        """
        try:
            if temp_md.exists():
                temp_md.unlink()
        except Exception as e:
            logger.warning(f"Could not delete temporary markdown file: {e}")

    def _set_language_for_run(self, rPr) -> None:
        """
        Set language for a run element.

        Args:
            rPr: Run properties element
        """
        lang_elements = rPr.findall(qn('w:lang'))
        for lang in lang_elements:
            rPr.remove(lang)

        lang = OxmlElement('w:lang')
        lang.set(qn('w:val'), self.config.language)
        lang.set(qn('w:eastAsia'), self.config.language)
        lang.set(qn('w:bidi'), self.config.language)
        rPr.append(lang)

    def _setup_footers(self, doc: Document) -> None:
        """
        Configure document footers with custom text and page numbers.

        Creates different footers for odd and even pages with page numbering.
        Odd pages: Right-aligned with "footer_text | Page X"
        Even pages: Left-aligned with "Page X | footer_text"

        Args:
            doc: Document object to modify
        """
        for section in doc.sections:
            # Configure section for different odd/even pages
            sectPr = section._sectPr
            if not sectPr.find(qn('w:titlePg')):
                titlePg = OxmlElement('w:titlePg')
                sectPr.append(titlePg)
            if not sectPr.find(qn('w:evenAndOddHeaders')):
                evenAndOddHeaders = OxmlElement('w:evenAndOddHeaders')
                sectPr.append(evenAndOddHeaders)

            # Configure headers/footers
            section.different_first_page_header_footer = True
            section.odd_and_even_pages_header_footer = True

            # Configure odd page footer (right pages)
            footer_odd = section.footer
            p_odd = footer_odd.paragraphs[0] if footer_odd.paragraphs else footer_odd.add_paragraph()
            p_odd.clear()
            p_odd.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p_odd.paragraph_format.space_before = Pt(12)

            # Add text and page number for odd pages
            run_odd = p_odd.add_run(f"{self.config.footer_text['odd']} | ")
            run_odd.font.name = self.config.font_name
            run_odd.font.size = Pt(DocumentConfig.DEFAULT_FOOTER_FONT_SIZE)
            self._add_page_number(p_odd)

            # Configure even page footer (left pages)
            footer_even = section.even_page_footer
            p_even = footer_even.paragraphs[0] if footer_even.paragraphs else footer_even.add_paragraph()
            p_even.clear()
            p_even.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p_even.paragraph_format.space_before = Pt(12)

            # Add page number and text for even pages
            self._add_page_number(p_even)
            run_even = p_even.add_run(f" | {self.config.footer_text['even']}")
            run_even.font.name = self.config.font_name
            run_even.font.size = Pt(DocumentConfig.DEFAULT_FOOTER_FONT_SIZE)

    def _add_page_number(self, paragraph) -> None:
        """
        Add a dynamic page number field to a paragraph.

        Args:
            paragraph: Paragraph object to add page number to
        """
        run = paragraph.add_run()

        # Begin field character
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar1)

        # Instruction text
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        run._r.append(instrText)

        # End field character
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar2)

    def _process_footnotes(self, doc: Document) -> None:
        """
        Configure footnote formatting and language settings.

        Args:
            doc: Document object to modify
        """
        try:
            # Configure footnote text style
            style = doc.styles['Footnote Text']
            style.font.name = self.config.font_name
            style.font.size = Pt(DocumentConfig.DEFAULT_FOOTNOTE_FONT_SIZE)
            style.paragraph_format.space_before = Pt(0)
            style.paragraph_format.space_after = Pt(0)
            style.paragraph_format.line_spacing = 1.0

            # Configure footnote reference style
            ref_style = doc.styles['Footnote Reference']
            ref_style.font.name = self.config.font_name
            ref_style.font.size = Pt(DocumentConfig.DEFAULT_FOOTNOTE_FONT_SIZE)
            ref_style.font.superscript = True

            # Apply language to footnote style
            style_element = style._element
            rPr = style_element.get_or_add_rPr()
            self._set_language_for_run(rPr)

            # Process each footnote if they exist
            if hasattr(doc, '_part') and hasattr(doc._part, '_footnotes_part') and doc._part._footnotes_part:
                footnotes = doc._part._footnotes_part.element.xpath('//w:footnote')
                for footnote in footnotes:
                    for p in footnote.xpath('.//w:p'):
                        pPr = p.get_or_add_pPr()

                        spacing = parse_xml(f'<w:spacing {qn("w")}:before="0" {qn("w")}:after="0" {qn("w")}:line="240" {qn("w")}:lineRule="auto"/>')
                        existing_spacing = pPr.find(qn('w:spacing'))
                        if existing_spacing is not None:
                            pPr.remove(existing_spacing)
                        pPr.append(spacing)

                        for r in p.xpath('.//w:r'):
                            rPr = r.get_or_add_rPr()
                            self._set_language_for_run(rPr)

        except KeyError as e:
            logger.warning(f"Footnote style not found: {e}")
        except Exception as e:
            logger.warning(f"Warning while processing footnotes: {str(e)}")

    def _post_process_document(self, doc_path: Path, document_title: str,
                               image_refs: List[dict], work_dir: Path) -> None:
        """
        Post-process the Word document with all formatting requirements.

        This method applies all styling, handles title formatting (including centering),
        inserts images, and processes all document elements.

        Args:
            doc_path: Path to the document file
            document_title: Main document title
            image_refs: List of image references to insert
            work_dir: Working directory containing img folder
        """
        doc = Document(doc_path)

        # Apply basic styles
        self._apply_global_styles(doc)

        # Find and style the title (first occurrence)
        self._style_main_title(doc, document_title)

        # Process all remaining paragraphs and insert images
        self._process_paragraphs_and_images(doc, image_refs, work_dir)

        # Process tables
        self._process_tables(doc)

        # Process footnotes
        self._process_footnotes(doc)

        # Setup footers
        self._setup_footers(doc)

        # Save the changes
        try:
            doc.save(doc_path)
        except Exception as e:
            raise IOError(f"Cannot save document: {e}")

    def _style_main_title(self, doc: Document, document_title: str) -> None:
        """
        Style the main title and optionally center it.
        Also styles the author and date paragraphs that follow.

        Args:
            doc: Document object
            document_title: Title to style
        """
        title_found = False
        title_index = -1
        
        for i, para in enumerate(doc.paragraphs):
            if para.text == document_title and not title_found:
                # Style as Title
                para.style = doc.styles['Title']
                para.clear()
                run = para.add_run(document_title)
                run.font.name = self.config.font_name
                run.font.size = Pt(DocumentConfig.DEFAULT_TITLE_SIZE)
                run.font.bold = True
                
                # Center the title if configured
                if self.config.center_title:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                title_found = True
                title_index = i
                
            # Style author and date paragraphs (next 2 paragraphs after title)
            elif title_found and i in [title_index + 1, title_index + 2]:
                # Keep existing text but reformat
                text = para.text
                para.clear()
                run = para.add_run(text)
                run.font.name = self.config.font_name
                run.font.size = Pt(self.config.base_font_size)
                run.font.bold = False
                
                # Center if title is centered
                if self.config.center_title:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
            elif title_found and i > title_index + 2:
                # Stop after processing title, author, and date
                break

    def _process_paragraphs_and_images(self, doc: Document, image_refs: List[dict],
                                      work_dir: Path) -> None:
        """
        Process all paragraphs in the document and insert images.

        Args:
            doc: Document object to modify
            image_refs: List of image references to insert
            work_dir: Working directory containing img folder
        """
        img_dir = work_dir / "img"

        for para in doc.paragraphs:
            # Check for image placeholder
            if '[IMAGE_PLACEHOLDER]' in para.text and image_refs:
                self._insert_single_image(para, image_refs, img_dir)
                continue

            # Skip paragraphs with specific styles
            if para.style.name in ['Title', 'Heading 1', 'Heading 2', 'Heading 3']:
                continue

            # Process paragraphs with 'Normal' style or no specific style
            if not para.style or para.style.name == 'Normal':
                self._format_normal_paragraph(para)

    def _insert_single_image(self, para, image_refs: List[dict], img_dir: Path) -> None:
        """
        Insert a single image into a paragraph.

        Args:
            para: Paragraph to insert image into
            image_refs: List of image references
            img_dir: Image directory path
        """
        img_ref = image_refs.pop(0)
        image_path = img_dir / img_ref['path']

        if image_path.suffix.lower() not in self.SUPPORTED_IMAGE_EXTENSIONS:
            logger.warning(f"Unsupported image format: {image_path.suffix}")

        logger.info(f"Processing image: {image_path}")

        if image_path.exists():
            para.clear()
            run = para.add_run()
            try:
                picture = run.add_picture(str(image_path))

                # Set image size with aspect ratio preservation
                if picture.width > self.MAX_IMAGE_WIDTH:
                    aspect_ratio = picture.height / picture.width
                    picture.width = self.MAX_IMAGE_WIDTH
                    picture.height = int(picture.width * aspect_ratio)

                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                logger.info(f"Added image: {img_ref['path']}")

            except Exception as e:
                logger.error(f"Error adding image {img_ref['path']}: {str(e)}")
                para.text = f"[Image: {img_ref['alt_text']}]"
        else:
            logger.warning(f"Image not found: {image_path}")
            para.text = f"[Image not found: {img_ref['alt_text']}]"

    def _format_normal_paragraph(self, para) -> None:
        """
        Apply formatting to a normal paragraph.

        Args:
            para: Paragraph to format
        """
        para.paragraph_format.space_before = Pt(DocumentConfig.DEFAULT_PARAGRAPH_SPACING)
        para.paragraph_format.space_after = Pt(DocumentConfig.DEFAULT_PARAGRAPH_SPACING)
        para.paragraph_format.line_spacing = self.config.line_spacing

        for run in para.runs:
            run.font.name = self.config.font_name
            run.font.size = Pt(self.config.base_font_size)

            # Set language for this run
            run_element = run._element
            rPr = run_element.get_or_add_rPr()
            self._set_language_for_run(rPr)

    def _apply_global_styles(self, doc: Document) -> None:
        """
        Apply global document styles including language and formatting.

        Args:
            doc: Document object to modify
        """
        self._set_document_language(doc)
        self._configure_standard_styles(doc)
        self._configure_heading_styles(doc)
        self._configure_section_properties(doc)

    def _set_document_language(self, doc: Document) -> None:
        """
        Set the document-wide language setting.

        Args:
            doc: Document object to modify
        """
        element = OxmlElement('w:lang')
        element.set(qn('w:val'), self.config.language)
        element.set(qn('w:eastAsia'), self.config.language)
        element.set(qn('w:bidi'), self.config.language)

        styles_element = doc.styles.element
        if styles_element.find(qn('w:docDefaults')) is None:
            doc_defaults = OxmlElement('w:docDefaults')
            styles_element.insert(0, doc_defaults)
        else:
            doc_defaults = styles_element.find(qn('w:docDefaults'))

        if doc_defaults.find(qn('w:rPrDefault')) is None:
            r_pr_default = OxmlElement('w:rPrDefault')
            doc_defaults.insert(0, r_pr_default)
        else:
            r_pr_default = doc_defaults.find(qn('w:rPrDefault'))

        if r_pr_default.find(qn('w:rPr')) is None:
            r_pr = OxmlElement('w:rPr')
            r_pr_default.insert(0, r_pr)
        else:
            r_pr = r_pr_default.find(qn('w:rPr'))

        existing_lang = r_pr.find(qn('w:lang'))
        if existing_lang is not None:
            r_pr.remove(existing_lang)

        r_pr.append(element)

    def _configure_standard_styles(self, doc: Document) -> None:
        """
        Configure standard document styles (Normal, Title).

        Args:
            doc: Document object to modify
        """
        styles_to_configure = {
            'Normal': {
                'font_size': self.config.base_font_size,
                'font_name': self.config.font_name,
                'bold': False,
                'space_before': 0,
                'space_after': 0,
                'line_spacing': self.config.line_spacing
            },
            'Title': {
                'font_size': DocumentConfig.DEFAULT_TITLE_SIZE,
                'font_name': self.config.font_name,
                'bold': True,
                'space_before': 12,
                'space_after': 12,
                'line_spacing': 1.0
            }
        }

        self._apply_style_configurations(doc, styles_to_configure)

    def _configure_heading_styles(self, doc: Document) -> None:
        """
        Configure heading styles (Heading 1-3).

        Args:
            doc: Document object to modify
        """
        heading_styles = {
            'Heading 1': {
                'font_size': DocumentConfig.DEFAULT_HEADING_1_SIZE,
                'font_name': self.config.font_name,
                'bold': True,
                'space_before': 18,
                'space_after': 12,
                'line_spacing': 1.0,
                'color': self.config.heading_colors.get(1, (0, 0, 0)),
                'italic': False
            },
            'Heading 2': {
                'font_size': DocumentConfig.DEFAULT_HEADING_2_SIZE,
                'font_name': self.config.font_name,
                'bold': True,
                'space_before': 16,
                'space_after': 10,
                'line_spacing': 1.0,
                'color': self.config.heading_colors.get(2, (0, 0, 0)),
                'italic': False
            },
            'Heading 3': {
                'font_size': DocumentConfig.DEFAULT_HEADING_3_SIZE,
                'font_name': self.config.font_name,
                'bold': True,
                'space_before': 14,
                'space_after': 8,
                'line_spacing': 1.0,
                'color': self.config.heading_colors.get(3, (0, 0, 0)),
                'italic': True
            }
        }

        self._apply_style_configurations(doc, heading_styles, is_heading=True)

    def _apply_style_configurations(self, doc: Document, styles_config: dict,
                                   is_heading: bool = False) -> None:
        """
        Apply style configurations to document.

        Args:
            doc: Document object to modify
            styles_config: Dictionary of style configurations
            is_heading: Whether these are heading styles
        """
        for style_name, config in styles_config.items():
            try:
                style = doc.styles[style_name]

                # Font configuration
                style.font.name = config['font_name']
                style.font.size = Pt(config['font_size'])
                style.font.bold = config['bold']
                if config.get('italic'):
                    style.font.italic = True

                # Color (if specified)
                if 'color' in config:
                    style.font.color.rgb = RGBColor(*config['color'])

                # Paragraph formatting
                style.paragraph_format.space_before = Pt(config['space_before'])
                style.paragraph_format.space_after = Pt(config['space_after'])
                style.paragraph_format.line_spacing = config['line_spacing']

                # For heading styles, ensure they're not linked to other styles
                if is_heading:
                    if hasattr(style, 'base_style'):
                        style.base_style = None

                    # Remove any existing numbering
                    if hasattr(style._element, 'pPr'):
                        pPr = style._element.pPr
                        if pPr is not None:
                            numPr = pPr.find(qn('w:numPr'))
                            if numPr is not None:
                                pPr.remove(numPr)

            except KeyError:
                logger.warning(f"Style '{style_name}' not found")

    def _configure_section_properties(self, doc: Document) -> None:
        """
        Configure section properties (page size, margins).

        Args:
            doc: Document object to modify
        """
        for section in doc.sections:
            # Set page size
            if self.config.paper_size == PaperSize.LETTER:
                section.page_width = Inches(8.5)
                section.page_height = Inches(11)
            elif self.config.paper_size == PaperSize.LEGAL:
                section.page_width = Inches(8.5)
                section.page_height = Inches(14)
            else:  # A4
                section.page_width = Inches(8.27)
                section.page_height = Inches(11.69)

            # Set margins
            section.top_margin = Cm(self.config.margins[0])
            section.right_margin = Cm(self.config.margins[1])
            section.bottom_margin = Cm(self.config.margins[2])
            section.left_margin = Cm(self.config.margins[3])

    def _process_tables(self, doc: Document) -> None:
        """
        Process all tables in the document with formatting and borders.

        Args:
            doc: Document object to modify
        """
        for table in doc.tables:
            # Process header row (first row)
            if table.rows:
                for cell in table.rows[0].cells:
                    self._format_table_cell(cell, is_header=True)

            # Process all rows
            for row in table.rows:
                for cell in row.cells:
                    self._format_table_cell(cell, is_header=False)
                    self._set_cell_borders(cell)

    def _format_table_cell(self, cell: _Cell, is_header: bool = False) -> None:
        """
        Format a table cell with appropriate styling.

        Args:
            cell: Cell to format
            is_header: Whether this is a header cell
        """
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.name = self.config.font_name
                run.font.size = Pt(DocumentConfig.DEFAULT_TABLE_FONT_SIZE)
                if is_header:
                    run.font.bold = True

                # Set language for this run
                run_element = run._element
                rPr = run_element.get_or_add_rPr()
                self._set_language_for_run(rPr)

            # Ensure paragraph has at least one run
            if not para.runs:
                run = para.add_run()
                run.font.name = self.config.font_name
                run.font.size = Pt(DocumentConfig.DEFAULT_TABLE_FONT_SIZE)
                if is_header:
                    run.font.bold = True

            # Set paragraph properties
            para.paragraph_format.space_before = Pt(DocumentConfig.DEFAULT_TABLE_CELL_SPACING)
            para.paragraph_format.space_after = Pt(DocumentConfig.DEFAULT_TABLE_CELL_SPACING)
            para.paragraph_format.line_spacing = 1.0

    def _set_cell_borders(self, cell: _Cell) -> None:
        """
        Add borders to a table cell if not already present.

        Args:
            cell: Cell object to add borders to
        """
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        # Check if borders already exist
        existing_borders = tcPr.find(qn('w:tcBorders'))
        if existing_borders is not None:
            return  # Borders already set

        # Create borders element
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

        # Add each border
        for border in ['top', 'left', 'bottom', 'right']:
            edge = OxmlElement(f'w:{border}')
            edge.set(qn('w:val'), 'single')
            edge.set(qn('w:sz'), '4')
            edge.set(qn('w:space'), '0')
            edge.set(qn('w:color'), 'auto')
            tcBorders.append(edge)


def main():
    """Example usage demonstrating the converter capabilities."""
    # Example configuration for a professional report
    config = DocumentConfig.create_report_style(
        author="Your Name",
        date="2025-01-31",
        language="en-US",
        generate_toc=True,
        paper_size=PaperSize.LETTER,
        font_name="Arial",
        base_font_size=12,
        heading_colors={
            1: (37, 150, 190),
            2: (37, 150, 190),
            3: (37, 150, 190)
        },
        footer_text={
            "odd": "Document Title",
            "even": "Author Name"
        },
        center_title=True
    )

    # Create converter and process file
    converter = MarkdownToDocxConverter(config, verbose=True)

    try:
        converter.convert(
            input_file="/Users/alain/Downloads/guide.md",
            output_file="/Users/alain/Downloads/guide.docx",
            working_dir="/Users/alain/Downloads"
        )
    except Exception as e:
        logger.error(f"Error during conversion: {str(e)}")


if __name__ == "__main__":
    main()